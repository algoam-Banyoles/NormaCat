"""
norm_monitor.py — Monitor de canvis normatius.

Compara snapshots dels catalegs per detectar:
  - Normes noves
  - Canvis d'estat (VIGENT -> DEROGADA, etc.)
  - Normes eliminades/retirades
  - Canvis de titol o metadades

Us:
    python tools/norm_monitor.py snapshot     # Desa snapshot actual
    python tools/norm_monitor.py check        # Compara amb ultim snapshot
    python tools/norm_monitor.py check --docx # Idem + genera informe DOCX
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import config

SNAPSHOTS_DIR = config.PROJECT_ROOT / "data" / "snapshots"
CHANGES_DIR = config.PROJECT_ROOT / "data" / "changes"


def _load_catalog(path: str) -> list[dict]:
    """Carrega un cataleg JSON (llista o dict amb llista dins)."""
    if not os.path.exists(path):
        return []
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, list):
        return data
    if isinstance(data, dict):
        for key in ("documents", "normes", "items", "resultats",
                     "legislacio", "llista", "data"):
            if key in data and isinstance(data[key], list):
                return data[key]
    return []


def _entry_key(entry: dict) -> str:
    """Genera una clau unica per a cada entrada."""
    codi = (entry.get("codi") or entry.get("id") or
            entry.get("boe_id") or entry.get("celex") or "").strip()
    titol = (entry.get("titol") or entry.get("title") or
             entry.get("text") or entry.get("nom") or "").strip()
    return codi if codi else titol[:100]


def _entry_status(entry: dict) -> str:
    """Extreu l'estat d'una entrada."""
    raw = (entry.get("estat") or entry.get("status") or
           entry.get("estat_legal") or "vigent").lower()
    if "derog" in raw or "anulad" in raw or "historic" in raw:
        return "DEROGADA"
    elif "referencia" in raw:
        return "REFERENCIA"
    elif "retirat" in raw:
        return "RETIRAT"
    return "VIGENT"


def take_snapshot() -> str:
    """Desa un snapshot de tots els catalegs actuals."""
    SNAPSHOTS_DIR.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    snapshot = {
        "timestamp": datetime.now().isoformat(),
        "sources": {},
    }

    total = 0
    for source_key, source_info in config.SOURCES.items():
        catalog_path = str(config.PROJECT_ROOT / source_info["catalog"])
        entries = _load_catalog(catalog_path)

        indexed = {}
        for entry in entries:
            key = _entry_key(entry)
            if not key:
                continue
            indexed[key] = {
                "status": _entry_status(entry),
                "title": (entry.get("titol") or entry.get("title") or
                         entry.get("text") or "").strip()[:200],
                "substituted_by": (entry.get("derogada_per") or
                                   entry.get("substituted_by") or
                                   entry.get("deroga") or "").strip(),
            }

        snapshot["sources"][source_key] = indexed
        total += len(indexed)
        print(f"  {source_info['label']:20s} {len(indexed):>6d} entrades")

    snapshot["total_entries"] = total

    snap_path = SNAPSHOTS_DIR / f"snapshot_{timestamp}.json"
    with open(snap_path, "w", encoding="utf-8") as f:
        json.dump(snapshot, f, ensure_ascii=False, indent=2)

    latest_path = SNAPSHOTS_DIR / "snapshot_latest.json"
    with open(latest_path, "w", encoding="utf-8") as f:
        json.dump(snapshot, f, ensure_ascii=False, indent=2)

    print(f"\n  Snapshot desat: {snap_path}")
    print(f"  Total: {total} entrades de {len(config.SOURCES)} fonts")

    return str(snap_path)


def check_changes(generate_docx=False) -> dict:
    """Compara catalegs actuals amb l'ultim snapshot."""
    latest_path = SNAPSHOTS_DIR / "snapshot_latest.json"
    if not latest_path.exists():
        print("  ERROR: No hi ha cap snapshot anterior.")
        print("  Executa primer: python tools/norm_monitor.py snapshot")
        return {}

    with open(latest_path, encoding="utf-8") as f:
        old_snap = json.load(f)

    old_timestamp = old_snap.get("timestamp", "desconegut")
    print(f"  Comparant amb snapshot de: {old_timestamp}")
    print()

    changes = {
        "old_timestamp": old_timestamp,
        "new_timestamp": datetime.now().isoformat(),
        "noves": [],
        "eliminades": [],
        "canvi_estat": [],
        "canvi_titol": [],
    }

    for source_key, source_info in config.SOURCES.items():
        label = source_info["label"]
        catalog_path = str(config.PROJECT_ROOT / source_info["catalog"])

        current_entries = _load_catalog(catalog_path)
        current = {}
        for entry in current_entries:
            key = _entry_key(entry)
            if not key:
                continue
            current[key] = {
                "status": _entry_status(entry),
                "title": (entry.get("titol") or entry.get("title") or
                         entry.get("text") or "").strip()[:200],
                "substituted_by": (entry.get("derogada_per") or
                                   entry.get("substituted_by") or
                                   entry.get("deroga") or "").strip(),
            }

        old = old_snap.get("sources", {}).get(source_key, {})

        for key in current:
            if key not in old:
                changes["noves"].append({
                    "source": label,
                    "key": key,
                    "status": current[key]["status"],
                    "title": current[key]["title"],
                })

        for key in old:
            if key not in current:
                changes["eliminades"].append({
                    "source": label,
                    "key": key,
                    "old_status": old[key]["status"],
                    "title": old[key]["title"],
                })

        for key in current:
            if key in old:
                if current[key]["status"] != old[key]["status"]:
                    changes["canvi_estat"].append({
                        "source": label,
                        "key": key,
                        "old_status": old[key]["status"],
                        "new_status": current[key]["status"],
                        "title": current[key]["title"],
                        "substituted_by": current[key].get("substituted_by", ""),
                    })
                elif current[key]["title"] != old[key]["title"]:
                    changes["canvi_titol"].append({
                        "source": label,
                        "key": key,
                        "old_title": old[key]["title"][:80],
                        "new_title": current[key]["title"][:80],
                    })

    # Mostrar resultats
    n_noves = len(changes["noves"])
    n_elim = len(changes["eliminades"])
    n_estat = len(changes["canvi_estat"])
    n_titol = len(changes["canvi_titol"])

    print(f"  {'='*50}")
    print(f"  RESULTATS DEL MONITOR DE CANVIS")
    print(f"  {'='*50}")
    print(f"  Noves entrades:    {n_noves}")
    print(f"  Eliminades:        {n_elim}")
    print(f"  Canvis d'estat:    {n_estat}")
    print(f"  Canvis de titol:   {n_titol}")
    print()

    if n_estat > 0:
        print(f"  CANVIS D'ESTAT (ATENCIO):")
        for c in changes["canvi_estat"]:
            print(f"    [{c['source']}] {c['key']}")
            print(f"      {c['old_status']} -> {c['new_status']}")
            if c.get("substituted_by"):
                print(f"      Substituida per: {c['substituted_by']}")
        print()

    if 0 < n_noves <= 20:
        print(f"  Noves entrades:")
        for c in changes["noves"]:
            print(f"    [{c['source']}] {c['key']} ({c['status']})")
        print()
    elif n_noves > 20:
        print(f"  Noves entrades: {n_noves} (massa per mostrar)")
        print()

    if 0 < n_elim <= 20:
        print(f"  Eliminades:")
        for c in changes["eliminades"]:
            print(f"    [{c['source']}] {c['key']}")
        print()

    # Desar informe JSON
    CHANGES_DIR.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    changes_path = CHANGES_DIR / f"canvis_{ts}.json"
    with open(changes_path, "w", encoding="utf-8") as f:
        json.dump(changes, f, ensure_ascii=False, indent=2)
    print(f"  Informe desat: {changes_path}")

    # Opcional: DOCX
    if generate_docx and (n_estat > 0 or n_noves > 0 or n_elim > 0):
        try:
            _generate_changes_docx(changes, ts)
        except ImportError:
            print("  [WARN] python-docx no disponible, DOCX no generat")

    return changes


def _generate_changes_docx(changes: dict, timestamp: str):
    """Genera un DOCX amb el resum de canvis."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    doc.add_heading("Monitor de Canvis Normatius -- NormaCat", level=1)
    doc.add_paragraph(f"Snapshot anterior: {changes.get('old_timestamp', '?')}")
    doc.add_paragraph(f"Data actual: {changes.get('new_timestamp', '?')}")
    doc.add_paragraph("")

    doc.add_heading("Resum", level=2)
    p = doc.add_paragraph()
    p.add_run(f"Noves entrades: {len(changes['noves'])}\n")
    p.add_run(f"Eliminades: {len(changes['eliminades'])}\n")
    run = p.add_run(f"Canvis d'estat: {len(changes['canvi_estat'])}\n")
    if changes["canvi_estat"]:
        run.font.color.rgb = RGBColor(220, 38, 38)
        run.bold = True

    if changes["canvi_estat"]:
        doc.add_heading("Canvis d'estat", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for i, txt in enumerate(["Norma", "Font", "Antic", "Nou"]):
            table.rows[0].cells[i].text = txt
            for par in table.rows[0].cells[i].paragraphs:
                for run in par.runs:
                    run.bold = True
        for c in changes["canvi_estat"]:
            row = table.add_row().cells
            row[0].text = c["key"]
            row[1].text = c["source"]
            row[2].text = c["old_status"]
            row[3].text = c["new_status"]

    if changes["noves"]:
        doc.add_heading("Noves entrades", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        for i, txt in enumerate(["Norma", "Font", "Estat"]):
            table.rows[0].cells[i].text = txt
            for par in table.rows[0].cells[i].paragraphs:
                for run in par.runs:
                    run.bold = True
        for c in changes["noves"][:50]:
            row = table.add_row().cells
            row[0].text = c["key"]
            row[1].text = c["source"]
            row[2].text = c["status"]

    docx_path = CHANGES_DIR / f"canvis_{timestamp}.docx"
    doc.save(str(docx_path))
    print(f"  Informe DOCX: {docx_path}")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    parser = argparse.ArgumentParser(
        description="Monitor de canvis normatius NormaCat"
    )
    parser.add_argument("command", choices=["snapshot", "check"],
                        help="'snapshot' = desa estat actual; 'check' = compara amb anterior")
    parser.add_argument("--docx", action="store_true",
                        help="Genera informe DOCX a mes del JSON")
    args = parser.parse_args()

    print("=" * 50)
    print("  NormaCat -- Monitor de Canvis Normatius")
    print("=" * 50)
    print()

    if args.command == "snapshot":
        take_snapshot()
    elif args.command == "check":
        check_changes(generate_docx=args.docx)


if __name__ == "__main__":
    main()
