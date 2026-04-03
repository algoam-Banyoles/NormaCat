"""
report_generator.py — Genera informes de conformitat normativa.

Input:  PDF (bytes o path) d'un annex normatiu de projecte
Output: DOCX amb l'analisi de totes les referencies trobades
"""

import io
import os
import re
import sys
from datetime import datetime
from pathlib import Path

import fitz  # PyMuPDF

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from search.norm_index import NormIndex
from search.norm_resolver import resolve as resolve_ref
import config

# ── Patrons d'extraccio de referencies ──────────────────────────────────────────

_REF_PATTERNS = [
    # RD, Llei, Decret, Ordre amb numero/any
    re.compile(r"\b(?:Reial\s+Decret|Real\s+Decreto|R\.?D\.?)\s*(?:Legisl\w+\s+)?(?:n[uu]m\.?\s*)?\d+\s*/\s*\d{2,4}\b", re.I),
    re.compile(r"\b(?:Llei|Ley)\s+(?:Organ\w+\s+)?(?:n[uu]m\.?\s*)?\d+\s*/\s*\d{2,4}\b", re.I),
    re.compile(r"(?<!\bReial\s)(?<!\bReal\s)\b(?:Decret|Decreto)\s+(?:n[uu]m\.?\s*)?\d+\s*/\s*\d{2,4}\b", re.I),
    re.compile(r"\b(?:Ordre|Orden)\s+[A-Z]+(?:/[A-Z]+)?/\d+/\d{4}\b", re.I),
    # UNE, ISO, EN
    re.compile(r"\bUNE(?:-EN)?(?:-ISO)?(?:/IEC)?\s+\d+(?:[-:/]\d+)*(?::\d{4})?\b", re.I),
    re.compile(r"(?<!\bUNE[-\s]EN\s)(?<!\bUNE\s)\bISO[/ ]?\d+(?:[-:/]\d+)*(?::\d{4})?\b", re.I),
    re.compile(r"(?<!UNE[-\s])\bEN\s+\d+(?:[-:/]\d+)*(?::\d{4})?\b", re.I),
    # Directives i Reglaments UE
    re.compile(r"\bDirecti(?:va|ve)\s+\(?\s*(?:UE|CE|CEE)\s*\)?\s*(?:n[uu]m\.?\s*)?\d+\s*/\s*\d{2,4}\b", re.I),
    re.compile(r"\bDirecti(?:va|ve)\s+\d+/\d+/(?:CE|UE|EU|CEE)\b", re.I),
    re.compile(r"\bReglament(?:o)?\s+\(?UE\)?\s+(?:n[uu]m\.?\s*)?\d+\s*/\s*\d{4}\b", re.I),
    # Normes amb nom
    re.compile(r"\b(?:EHE[-\s]?\d{2}|IAP[-\s]?\d{2}|NCSE[-\s]?\d{2})\b", re.I),
    re.compile(r"\b(?:PG[-\s]?3|REBT|RITE|RIPCI|EAE)\b", re.I),
    re.compile(r"\bCTE\s+DB[-\s]?[A-Z]{2,4}(?:[-\s]?\d)?\b", re.I),
    # NTE
    re.compile(r"\bNTE[-\s]?[A-Z]{2,4}[-\s]?\d{0,3}\b", re.I),
]


def extract_text_from_pdf(pdf_input) -> str:
    """Extreu text d'un PDF (bytes o path)."""
    if isinstance(pdf_input, (str, Path)):
        doc = fitz.open(str(pdf_input))
    else:
        doc = fitz.open(stream=pdf_input, filetype="pdf")

    pages_text = []
    for page in doc:
        text = page.get_text("text")
        if text and text.strip():
            pages_text.append(text)
    doc.close()
    return "\n\n".join(pages_text)


def extract_references(text: str) -> list[str]:
    """Extreu totes les referencies normatives uniques del text."""
    found = set()
    for pattern in _REF_PATTERNS:
        for match in pattern.finditer(text):
            ref = match.group(0).strip()
            if len(ref) > 3:
                found.add(ref)
    return sorted(found)


def validate_references(refs: list[str], norm_index: NormIndex) -> list[dict]:
    """Valida cada referencia contra el NormIndex."""
    results = []
    for ref in refs:
        resolved = resolve_ref(ref)
        lookup = norm_index.lookup(ref)

        status = "NO_TROBADA"
        if lookup:
            status = lookup.get("status", "PENDENT")

        results.append({
            "raw": ref,
            "resolved": resolved,
            "status": status,
            "source": lookup.get("source") if lookup else None,
            "title": lookup.get("title") if lookup else None,
            "substituted_by": lookup.get("substituted_by") if lookup else None,
            "fuzzy": lookup.get("fuzzy", False) if lookup else False,
        })

    order = {"DEROGADA": 0, "NO_TROBADA": 1, "PENDENT": 2,
             "REFERENCIA": 3, "VIGENT": 4}
    results.sort(key=lambda r: order.get(r["status"], 5))

    return results


def generate_docx_report(
    results: list[dict],
    project_name: str = "",
    pdf_filename: str = "",
) -> bytes:
    """Genera l'informe DOCX i retorna els bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Estils
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Capcalera
    h = doc.add_heading("Informe de Conformitat Normativa", level=1)
    h.runs[0].font.color.rgb = RGBColor(26, 58, 95)

    doc.add_paragraph(f"Projecte: {project_name or '(no especificat)'}")
    doc.add_paragraph(f"Fitxer analitzat: {pdf_filename or '(no especificat)'}")
    doc.add_paragraph(f"Data d'analisi: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("Generat per: NormaCat v1.0 -- Servei de Supervisio de Projectes (DGIM)")
    doc.add_paragraph("")

    # Comptadors
    n_total = len(results)
    n_vigent = sum(1 for r in results if r["status"] == "VIGENT")
    n_derogada = sum(1 for r in results if r["status"] == "DEROGADA")
    n_referencia = sum(1 for r in results if r["status"] == "REFERENCIA")
    n_pendent = sum(1 for r in results if r["status"] in ("PENDENT", "NO_TROBADA"))

    doc.add_heading("Resum", level=2)
    p = doc.add_paragraph()
    run = p.add_run(f"Total de referencies analitzades: {n_total}\n")
    run.bold = True
    p.add_run(f"  Vigents: {n_vigent}\n")
    run_d = p.add_run(f"  Derogades: {n_derogada}\n")
    if n_derogada > 0:
        run_d.font.color.rgb = RGBColor(220, 38, 38)
        run_d.bold = True
    p.add_run(f"  Referencia: {n_referencia}\n")
    p.add_run(f"  Pendents de verificacio: {n_pendent}\n")

    # ALERTES: Normes derogades
    derogades = [r for r in results if r["status"] == "DEROGADA"]
    if derogades:
        doc.add_heading("NORMES DEROGADES", level=2)
        doc.add_paragraph(
            "Les seguents normes citades al projecte estan DEROGADES. "
            "Cal substituir-les per la normativa vigent indicada."
        )

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        for i, text in enumerate(["Norma citada", "Estat", "Substituida per"]):
            hdr[i].text = text
            for par in hdr[i].paragraphs:
                for run in par.runs:
                    run.bold = True

        for r in derogades:
            row = table.add_row().cells
            row[0].text = r["raw"]
            row[1].text = "DEROGADA"
            row[2].text = r["substituted_by"] or "(consultar manualment)"
            for par in row[1].paragraphs:
                for run in par.runs:
                    run.font.color.rgb = RGBColor(220, 38, 38)
                    run.bold = True

        doc.add_paragraph("")

    # ATENCIO: No trobades
    no_trobades = [r for r in results if r["status"] in ("PENDENT", "NO_TROBADA")]
    if no_trobades:
        doc.add_heading("NORMES PENDENTS DE VERIFICACIO", level=2)
        doc.add_paragraph(
            "Les seguents normes no s'han pogut verificar automaticament. "
            "Cal comprovar-ne la vigencia manualment."
        )

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        hdr[0].text = "Norma citada"
        hdr[1].text = "Observacions"
        for cell in hdr:
            for par in cell.paragraphs:
                for run in par.runs:
                    run.bold = True

        for r in no_trobades:
            row = table.add_row().cells
            row[0].text = r["raw"]
            row[1].text = "Verificar vigencia manualment"
            if r.get("fuzzy"):
                row[1].text += " (coincidencia aproximada)"

        doc.add_paragraph("")

    # OK: Vigents
    vigents = [r for r in results if r["status"] == "VIGENT"]
    if vigents:
        doc.add_heading("NORMES VIGENTS", level=2)
        doc.add_paragraph(f"{len(vigents)} normes verificades com a vigents.")

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        for i, text in enumerate(["Norma", "Font", "Titol"]):
            hdr[i].text = text
            for par in hdr[i].paragraphs:
                for run in par.runs:
                    run.bold = True

        for r in vigents:
            row = table.add_row().cells
            row[0].text = r["raw"]
            row[1].text = r["source"] or "--"
            row[2].text = (r["title"] or "--")[:80]

    # Generar bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def analyze_pdf(pdf_input, project_name="", pdf_filename="") -> dict:
    """Pipeline complet: PDF -> extraccio -> validacio -> informe."""
    idx = NormIndex(str(config.PROJECT_ROOT))

    text = extract_text_from_pdf(pdf_input)
    refs = extract_references(text)
    results = validate_references(refs, idx)
    report_bytes = generate_docx_report(results, project_name, pdf_filename)

    stats = {
        "total": len(results),
        "vigent": sum(1 for r in results if r["status"] == "VIGENT"),
        "derogada": sum(1 for r in results if r["status"] == "DEROGADA"),
        "pendent": sum(1 for r in results if r["status"] in ("PENDENT", "NO_TROBADA")),
    }

    return {
        "text_length": len(text),
        "references": results,
        "report_bytes": report_bytes,
        "stats": stats,
    }
